import argparse
import csv
from pathlib import Path

# Try to import python-docx for DOCX export
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def parse_lhco(file_path):
    """
    Parse LHCO file into a list of records.
    Each event starts with a header line of 3 tokens (event header), followed by object lines.
    Returns a list of dicts with keys: event, idx, typ, eta, phi, pt, jmas, ntrk, btag, hadem.
    """
    records = []
    event_id = None
    with open(file_path, 'r') as f:
        for raw in f:
            tokens = raw.split()
            if not tokens:
                continue
            # Header line: e.g. '0   <event_id>   0'
            if len(tokens) == 3 and tokens[0].isdigit():
                event_id = tokens[1]
                continue
            # Object line: at least 9 meaningful columns (ignore dum1, dum2)
            if len(tokens) >= 9:
                rec = {
                    'event': event_id,
                    'idx': tokens[0],
                    'typ': tokens[1],
                    'eta': tokens[2],
                    'phi': tokens[3],
                    'pt': tokens[4],
                    'jmas': tokens[5],
                    'ntrk': tokens[6],
                    'btag': tokens[7],
                    'hadem': tokens[8],
                }
                records.append(rec)
    return records


def filter_records(records, particle=None, columns=None):
    """
    Filter by particle type (string) and select specific columns.
    Always includes 'event' in output.
    """
    out = []
    for rec in records:
        if particle is not None and rec['typ'] != str(particle):
            continue
        if columns:
            filtered = {'event': rec['event']}
            for c in columns:
                if c in rec:
                    filtered[c] = rec[c]
            out.append(filtered)
        else:
            out.append(rec.copy())
    return out


def write_csv(records, out_path):
    if not records:
        print("No records to write.")
        return
    with open(out_path, 'w', newline='') as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=list(records[0].keys()))
        writer.writeheader()
        writer.writerows(records)
    print(f"CSV written to {out_path}")


def write_docx(records, out_path):
    if not DOCX_AVAILABLE:
        print("Error: python-docx is not installed. Cannot write DOCX.")
        return
    if not records:
        print("No records to write.")
        return
    doc = Document()
    doc.add_heading('LHCO Data', level=1)
    cols = list(records[0].keys())
    table = doc.add_table(rows=1, cols=len(cols))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr_cells[i].text = col
    for rec in records:
        row_cells = table.add_row().cells
        for i, col in enumerate(cols):
            row_cells[i].text = str(rec[col])
    doc.save(out_path)
    print(f"DOCX written to {out_path}")


def main():
    parser = argparse.ArgumentParser(description='Convert LHCO to CSV or DOCX')
    parser.add_argument('input', help='Input LHCO file path')
    parser.add_argument('output', help='Output file path (.csv or .docx)')
    parser.add_argument('--format', choices=['csv','docx'], default='csv',
                        help='Output format')
    parser.add_argument('--particle', help='Filter by particle typ (e.g. 2)')
    parser.add_argument('--columns', nargs='+', help='Columns to include (e.g. pt eta phi)')
    args = parser.parse_args()

    infile = Path(args.input)
    if not infile.exists():
        print(f"Input file {infile} not found.", file=sys.stderr)
        sys.exit(1)

    recs = parse_lhco(infile)
    recs = filter_records(recs, particle=args.particle, columns=args.columns)

    if args.format == 'csv':
        write_csv(recs, args.output)
    else:
        write_docx(recs, args.output)

    if args.format == 'docx' and not DOCX_AVAILABLE:
        print("Reminder: install python-docx via 'pip install python-docx' to enable DOCX output.")

if __name__ == '__main__':
    main()

#output format
#python lhco_converter.py input.lhco output.csv
#python lhco_converter.py input.lhco output.docx --format docx
